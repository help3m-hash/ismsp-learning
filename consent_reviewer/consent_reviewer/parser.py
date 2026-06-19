"""문서 파서 (Dev2 담당).

다양한 형식의 동의서 파일을 읽어 판단 엔진(Claude)이 점검할 수 있는
`ParsedDocument` 형태로 변환한다.

핵심 설계
---------
Claude는 PDF와 이미지를 네이티브로 직접 읽을 수 있다(vision / PDF document block).
따라서 두 갈래로 처리한다.

- 텍스트 추출형(txt / docx / html / 텍스트 레이어가 있는 PDF):
  본문 텍스트를 추출해 `text`에 담는다. (`media_blocks`는 비움)
- 비텍스트형(스캔 PDF / 이미지):
  OCR을 직접 하지 않고 base64로 인코딩해 Claude content block(`media_blocks`)에 담는다.
  판단 엔진이 그대로 Claude API에 전달한다. (`text`는 None)
"""

from __future__ import annotations

import base64
from pathlib import Path

from .models import ParsedDocument

# ---- 지원 확장자 ---------------------------------------------------------

# 이미지 확장자 → Claude content block에 넣을 media_type 매핑.
_IMAGE_MEDIA_TYPES: dict[str, str] = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".gif": "image/gif",
}

SUPPORTED_EXTS: set[str] = {
    ".txt",
    ".docx",
    ".html",
    ".htm",
    ".pdf",
    *_IMAGE_MEDIA_TYPES.keys(),
}

# ---- PDF 텍스트/스캔 판별 임계값 -----------------------------------------

# pdfplumber로 추출한 텍스트에서 공백(스페이스/개행/탭)을 제외한 글자 수가
# 이 값 이상이면 "텍스트 레이어가 있는 PDF"(pdf_text)로 간주한다.
# 그 미만이면 스캔본/이미지 PDF로 보고 base64 document block(pdf_image)으로 폴백한다.
# - 너무 작게 잡으면 페이지 번호·머리글만 있는 스캔본을 텍스트로 오인한다.
# - 너무 크게 잡으면 짧은 동의서(반 페이지)를 스캔본으로 오인한다.
# 동의서는 보통 수백 자 이상이므로 80자를 보수적 기준으로 둔다.
PDF_TEXT_MIN_CHARS = 80


# ---- 공개 API ------------------------------------------------------------


def parse_document(path: str) -> ParsedDocument:
    """경로의 파일을 형식에 맞게 파싱해 `ParsedDocument`를 반환한다.

    Args:
        path: 파싱할 파일 경로.

    Returns:
        ParsedDocument: 텍스트 추출형이면 `text`, 비텍스트형이면 `media_blocks`.

    Raises:
        FileNotFoundError: 경로에 파일이 없을 때.
        ValueError: 지원하지 않는 확장자일 때.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"파일을 찾을 수 없습니다: {path}")
    if not p.is_file():
        raise FileNotFoundError(f"파일이 아닙니다(디렉터리 등): {path}")

    ext = p.suffix.lower()
    if ext not in SUPPORTED_EXTS:
        supported = ", ".join(sorted(SUPPORTED_EXTS))
        raise ValueError(
            f"지원하지 않는 확장자입니다: '{ext}'. 지원 형식: {supported}"
        )

    if ext == ".txt":
        return _parse_txt(p)
    if ext == ".docx":
        return _parse_docx(p)
    if ext in (".html", ".htm"):
        return _parse_html(p)
    if ext == ".pdf":
        return _parse_pdf(p)
    # 나머지는 모두 이미지 확장자.
    return _parse_image(p)


# ---- 형식별 파서 ---------------------------------------------------------


def _parse_txt(p: Path) -> ParsedDocument:
    """.txt: UTF-8 우선, 실패 시 cp949 폴백으로 읽는다."""
    raw = p.read_bytes()
    if not raw.strip():
        return ParsedDocument(
            source_path=str(p),
            file_type="txt",
            text="",
            parser_note="빈 파일입니다.",
        )

    note = ""
    text: str | None = None
    # UTF-8(BOM 포함) → cp949 순으로 시도.
    for enc in ("utf-8-sig", "cp949"):
        try:
            text = raw.decode(enc)
            if enc != "utf-8-sig":
                note = f"UTF-8 디코딩 실패로 {enc}로 읽었습니다."
            break
        except UnicodeDecodeError:
            continue

    if text is None:
        # 두 인코딩 모두 실패 — 손상/미지원 인코딩. 깨진 문자는 대체해 본문을 살린다.
        text = raw.decode("utf-8", errors="replace")
        note = "UTF-8/cp949 디코딩에 모두 실패해 일부 문자를 대체(replace)했습니다."

    return ParsedDocument(
        source_path=str(p),
        file_type="txt",
        text=text,
        parser_note=note,
    )


def _parse_docx(p: Path) -> ParsedDocument:
    """.docx: 문단 + 표 셀 텍스트까지 추출(동의서는 표 형태가 많음)."""
    try:
        import docx  # python-docx
    except ImportError as e:  # pragma: no cover - 의존성 누락 환경
        raise ImportError(
            "python-docx가 필요합니다. `pip install python-docx`로 설치하세요."
        ) from e

    try:
        document = docx.Document(str(p))
    except Exception as e:  # 손상 파일/잘못된 zip 등
        return ParsedDocument(
            source_path=str(p),
            file_type="docx",
            text="",
            parser_note=f"docx 열기에 실패했습니다(손상 가능): {e}",
        )

    parts: list[str] = []

    # 1) 본문 문단.
    for para in document.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)

    # 2) 표 셀 텍스트(동의서는 '수집항목/목적/보유기간' 등이 표로 들어가는 경우가 많음).
    #    셀은 탭으로, 행은 개행으로 구분해 표 구조를 어느 정도 보존한다.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))

    text = "\n".join(parts)
    note = "" if text.strip() else "추출된 텍스트가 없습니다(빈 문서이거나 이미지 기반일 수 있음)."

    return ParsedDocument(
        source_path=str(p),
        file_type="docx",
        text=text,
        parser_note=note,
    )


def _parse_html(p: Path) -> ParsedDocument:
    """.html/.htm: script/style 제거 후 텍스트 추출."""
    try:
        from bs4 import BeautifulSoup
    except ImportError as e:  # pragma: no cover
        raise ImportError(
            "beautifulsoup4가 필요합니다. `pip install beautifulsoup4`로 설치하세요."
        ) from e

    raw = p.read_bytes()
    if not raw.strip():
        return ParsedDocument(
            source_path=str(p),
            file_type="html",
            text="",
            parser_note="빈 파일입니다.",
        )

    # 인코딩은 bs4가 메타 태그 등을 보고 추정하도록 bytes를 그대로 넘긴다.
    soup = BeautifulSoup(raw, "html.parser")

    # 화면에 보이지 않는 요소 제거.
    for tag in soup(["script", "style", "noscript", "head", "meta", "link"]):
        tag.decompose()

    # separator로 블록 간 공백을 확보하고, 줄별로 빈 줄을 정리한다.
    text = soup.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines()]
    text = "\n".join(line for line in lines if line)

    note = "" if text.strip() else "추출된 텍스트가 없습니다."

    return ParsedDocument(
        source_path=str(p),
        file_type="html",
        text=text,
        parser_note=note,
    )


def _parse_pdf(p: Path) -> ParsedDocument:
    """.pdf: 텍스트 추출을 시도하고, 충분하면 pdf_text, 거의 없으면 pdf_image로 폴백."""
    try:
        import pdfplumber
    except ImportError as e:  # pragma: no cover
        raise ImportError(
            "pdfplumber가 필요합니다. `pip install pdfplumber`로 설치하세요."
        ) from e

    extracted = ""
    extract_error = ""
    try:
        with pdfplumber.open(str(p)) as pdf:
            page_texts: list[str] = []
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t.strip():
                    page_texts.append(t)
            extracted = "\n".join(page_texts)
    except Exception as e:
        # 텍스트 추출 자체가 실패해도 스캔본일 수 있으므로 이미지 모드로 폴백을 시도한다.
        extract_error = f"텍스트 추출 중 오류: {e}"

    # 공백 제외 글자 수로 텍스트 레이어 충분 여부 판정.
    non_ws_len = len("".join(extracted.split()))

    if non_ws_len >= PDF_TEXT_MIN_CHARS:
        return ParsedDocument(
            source_path=str(p),
            file_type="pdf_text",
            text=extracted,
            parser_note=f"텍스트 레이어 추출 성공(공백 제외 {non_ws_len}자).",
        )

    # 텍스트가 거의 없음 → 스캔본/이미지 PDF로 보고 base64 document block 생성.
    reason_bits = []
    if extract_error:
        reason_bits.append(extract_error)
    reason_bits.append(
        f"추출 텍스트가 임계값 미만(공백 제외 {non_ws_len}자 < {PDF_TEXT_MIN_CHARS}자)"
    )
    note = (
        "스캔본/이미지 PDF로 판단하여 PDF document block(base64)으로 전달합니다. "
        + "사유: "
        + "; ".join(reason_bits)
        + "."
    )

    b64 = base64.standard_b64encode(p.read_bytes()).decode("ascii")
    media_block = {
        "type": "document",
        "source": {
            "type": "base64",
            "media_type": "application/pdf",
            "data": b64,
        },
    }

    return ParsedDocument(
        source_path=str(p),
        file_type="pdf_image",
        text=None,
        media_blocks=[media_block],
        parser_note=note,
    )


def _parse_image(p: Path) -> ParsedDocument:
    """이미지(.png/.jpg/.jpeg/.webp/.gif): base64 image block 생성."""
    ext = p.suffix.lower()
    media_type = _IMAGE_MEDIA_TYPES[ext]

    raw = p.read_bytes()
    if not raw:
        return ParsedDocument(
            source_path=str(p),
            file_type="image",
            text=None,
            parser_note="빈 이미지 파일입니다.",
        )

    b64 = base64.standard_b64encode(raw).decode("ascii")
    media_block = {
        "type": "image",
        "source": {
            "type": "base64",
            "media_type": media_type,
            "data": b64,
        },
    }

    return ParsedDocument(
        source_path=str(p),
        file_type="image",
        text=None,
        media_blocks=[media_block],
        parser_note=f"이미지를 base64 image block({media_type})으로 전달합니다.",
    )
